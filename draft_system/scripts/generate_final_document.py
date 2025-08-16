#!/usr/bin/env python3
"""
Final Document Generator
Generates PDF and Word documents from HTML preview using standardized output folder
"""

import os
import sys
import yaml
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict

# Import our existing generators
from universal_preview_generator import UniversalPreviewGenerator

class FinalDocumentGenerator:
    """Generates final PDF and Word documents from HTML preview"""
    
    def __init__(self):
        """Initialize the final document generator"""
        self.script_dir = Path(__file__).parent.parent
        self.preview_dir = self.script_dir / "html_preview"
        self.output_dir = self.script_dir / "final_documents"
        self.preview_file = self.preview_dir / "draft-preview.html"
        
        # Ensure output directory exists
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_documents(self, document_folder: str, formats: list = None) -> Dict[str, str]:
        """
        Generate final documents in specified formats
        
        Args:
            document_folder: Path to document folder
            formats: List of formats ['pdf', 'docx', 'both'] (default: ['both'])
            
        Returns:
            Dictionary with generated file paths
        """
        if formats is None:
            formats = ['both']
        
        folder = Path(document_folder)
        if not folder.exists():
            raise FileNotFoundError(f"Document folder not found: {document_folder}")
        
        # Load metadata to get document info
        metadata = self._load_metadata(folder)
        document_info = metadata.get('document_info', {})
        document_type = document_info.get('type', 'document')
        document_name = document_info.get('name', folder.name)
        
        # Generate timestamp for unique filenames
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"{document_type}_{document_name}_{timestamp}"
        
        # First, generate HTML preview
        preview_generator = UniversalPreviewGenerator()
        preview_file = preview_generator.generate_preview(document_folder)
        print(f"HTML preview generated: {preview_file}")

        # Verify the preview file exists and has content
        if not Path(preview_file).exists():
            raise FileNotFoundError(f"Preview file not generated: {preview_file}")

        with open(preview_file, 'r', encoding='utf-8') as f:
            preview_content = f.read()

        if not preview_content.strip():
            raise ValueError("Preview file is empty")

        # Check if preview contains template syntax (indicates an error)
        if '{{' in preview_content and '}}' in preview_content:
            print("Warning: Preview contains unprocessed template syntax")
            # Continue anyway, but user should be aware
        
        generated_files = {}
        
        # Generate requested formats
        if 'pdf' in formats or 'both' in formats:
            pdf_file = self._generate_pdf(base_filename)
            if pdf_file:
                generated_files['pdf'] = pdf_file
        
        if 'docx' in formats or 'both' in formats:
            docx_file = self._generate_docx(base_filename)
            if docx_file:
                generated_files['docx'] = docx_file
        
        return generated_files
    
    def _load_metadata(self, folder: Path) -> Dict:
        """Load metadata.yml from document folder"""
        metadata_file = folder / "metadata.yml"
        if metadata_file.exists():
            with open(metadata_file, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        return {}
    
    def _generate_pdf(self, base_filename: str) -> Optional[str]:
        """Generate PDF from HTML preview using browser print-to-PDF"""
        pdf_filename = f"{base_filename}.pdf"
        pdf_path = self.output_dir / pdf_filename

        # Try browser-based PDF generation first (best formatting)
        if self._generate_pdf_browser(pdf_path):
            return str(pdf_path)

        # Fallback to weasyprint
        try:
            import weasyprint

            # Read HTML content
            with open(self.preview_file, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # Enhance CSS for PDF output
            enhanced_html = self._enhance_html_for_pdf(html_content)

            # Generate PDF
            weasyprint.HTML(string=enhanced_html).write_pdf(str(pdf_path))

            print(f"PDF generated (weasyprint): {pdf_path}")
            return str(pdf_path)

        except ImportError:
            print("Warning: weasyprint not installed. Trying alternative methods...")
            return self._generate_pdf_fallback(base_filename)
        except Exception as e:
            print(f"Error generating PDF with weasyprint: {e}")
            return self._generate_pdf_fallback(base_filename)

    def _generate_pdf_browser(self, pdf_path: Path) -> bool:
        """Generate PDF using browser print-to-PDF (best quality)"""
        try:
            import subprocess
            import platform

            # Get absolute path to HTML file
            html_path = self.preview_file.absolute()

            # Chrome/Chromium command for different platforms
            chrome_commands = []

            if platform.system() == "Darwin":  # macOS
                chrome_commands = [
                    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
                    "/Applications/Chromium.app/Contents/MacOS/Chromium"
                ]
            elif platform.system() == "Windows":
                chrome_commands = [
                    "chrome.exe",
                    "chromium.exe",
                    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"
                ]
            else:  # Linux
                chrome_commands = [
                    "google-chrome",
                    "chromium-browser",
                    "chromium"
                ]

            # Try each Chrome command
            for chrome_cmd in chrome_commands:
                try:
                    cmd = [
                        chrome_cmd,
                        "--headless",
                        "--disable-gpu",
                        "--print-to-pdf=" + str(pdf_path),
                        "--print-to-pdf-no-header",
                        "--no-margins",
                        f"file://{html_path}"
                    ]

                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

                    if result.returncode == 0 and pdf_path.exists():
                        print(f"PDF generated (browser): {pdf_path}")
                        return True

                except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError):
                    continue

            return False

        except Exception as e:
            print(f"Browser PDF generation failed: {e}")
            return False
    
    def _generate_pdf_fallback(self, base_filename: str) -> Optional[str]:
        """Fallback PDF generation using existing html_to_pdf.py"""
        try:
            import subprocess
            
            pdf_filename = f"{base_filename}.pdf"
            pdf_path = self.output_dir / pdf_filename
            
            # Use existing html_to_pdf script
            cmd = [
                sys.executable, 
                str(self.script_dir / "scripts" / "html_to_pdf.py"),
                "convert",
                str(self.preview_file),
                str(pdf_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0 and pdf_path.exists():
                print(f"PDF generated: {pdf_path}")
                return str(pdf_path)
            else:
                print(f"PDF generation failed: {result.stderr}")
                return None
                
        except Exception as e:
            print(f"Error in PDF fallback: {e}")
            return None
    
    def _generate_docx(self, base_filename: str) -> Optional[str]:
        """Generate Word document from HTML preview"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            import re
            from bs4 import BeautifulSoup
            
            docx_filename = f"{base_filename}.docx"
            docx_path = self.output_dir / docx_filename
            
            # Read and parse HTML content
            with open(self.preview_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create Word document
            doc = Document()
            
            # Set up document formatting
            self._setup_document_formatting(doc)
            
            # Extract and add content
            self._add_html_content_to_docx(doc, soup)
            
            # Save document
            doc.save(str(docx_path))
            
            print(f"Word document generated: {docx_path}")
            return str(docx_path)
            
        except ImportError:
            print("Warning: python-docx not installed. Install with: pip install python-docx beautifulsoup4")
            return None
        except Exception as e:
            print(f"Error generating Word document: {e}")
            return None
    
    def _enhance_html_for_pdf(self, html_content: str) -> str:
        """Enhance HTML with better CSS for PDF output"""
        # Add enhanced CSS for PDF
        enhanced_css = """
        <style>
        @page {
            size: letter;
            margin: 1in;
        }
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 2.0;
            color: #000;
            background: #fff;
        }
        .header {
            text-align: center;
            margin-bottom: 2em;
        }
        .court-name {
            font-weight: bold;
            font-size: 14pt;
        }
        .case-caption {
            margin: 2em 0;
            border-collapse: collapse;
            width: 100%;
        }
        .case-caption td {
            border: 1px solid #000;
            padding: 0.5em;
            vertical-align: top;
        }
        h1, h2, h3 {
            text-align: center;
            font-weight: bold;
            page-break-after: avoid;
        }
        h1 { font-size: 14pt; }
        h2 { font-size: 12pt; }
        h3 { font-size: 12pt; }
        p {
            text-align: justify;
            margin: 1em 0;
            orphans: 2;
            widows: 2;
        }
        .signature-block {
            margin-top: 3em;
            page-break-inside: avoid;
        }
        strong { font-weight: bold; }
        em { font-style: italic; }
        </style>
        """
        
        # Replace existing style tag
        import re
        html_content = re.sub(r'<style>.*?</style>', enhanced_css, html_content, flags=re.DOTALL)
        
        return html_content
    
    def _setup_document_formatting(self, doc):
        """Set up Word document formatting"""
        from docx.shared import Inches, Pt

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Set line spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 2.0
    
    def _add_html_content_to_docx(self, doc, soup):
        """Add HTML content to Word document"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        from docx.oxml.shared import OxmlElement, qn

        # First, add the caption section
        self._add_caption_to_docx(doc, soup)

        # Then add the document title
        self._add_document_title_to_docx(doc, soup)

        # Finally, add the main content
        content_div = soup.find('div', class_='content') or soup.find('body')

        if content_div:
            for element in content_div.find_all(['h1', 'h2', 'h3', 'p']):
                if element.name in ['h1', 'h2', 'h3']:
                    # Add heading
                    heading = doc.add_heading(element.get_text().strip(), level=int(element.name[1]))
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif element.name == 'p':
                    # Add paragraph
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # Process text with formatting
                    self._add_formatted_text(para, element)

    def _add_caption_to_docx(self, doc, soup):
        """Add caption section to Word document"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        caption_div = soup.find('div', class_='caption')
        if not caption_div:
            return

        # Add court name(s)
        court_divs = caption_div.find_all('div', class_='court')
        for court_div in court_divs:
            court_para = doc.add_paragraph()
            court_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = court_para.add_run(court_div.get_text().strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True

        # Add some space
        doc.add_paragraph()

        # Add caption table content
        caption_table = caption_div.find('table', class_='caption-grid')
        if caption_table:
            # Create a table in Word
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Get the cells
            cells = table.rows[0].cells

            # Left cell - parties
            parties_td = caption_table.find('td', class_='parties')
            if parties_td:
                parties_cell = cells[0]
                for div in parties_td.find_all('div'):
                    para = parties_cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    text = div.get_text().strip()
                    run = para.add_run(text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    if 'name' in div.get('class', []):
                        run.bold = True
                    elif 'role' in div.get('class', []):
                        run.italic = True

            # Middle cell - vertical line (empty)
            cells[1].text = ""

            # Right cell - case info
            meta_td = caption_table.find('td', class_='meta')
            if meta_td:
                meta_cell = cells[2]
                meta_cell.text = ""  # Clear default text
                for div in meta_td.find_all('div'):
                    para = meta_cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    text = div.get_text().strip()
                    run = para.add_run(text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    if 'label' in div.get('class', []):
                        run.bold = True

        # Add space after caption
        doc.add_paragraph()

    def _add_document_title_to_docx(self, doc, soup):
        """Add document title to Word document"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        title_div = soup.find('div', class_='doc-title')
        if title_div:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(title_div.get_text().strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True

            # Add space after title
            doc.add_paragraph()
    
    def _add_formatted_text(self, paragraph, element):
        """Add formatted text to Word paragraph"""
        from docx.shared import Pt
        
        for content in element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                else:
                    run = paragraph.add_run(content.get_text())
            else:
                # Plain text
                run = paragraph.add_run(str(content))
            
            # Set font
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

def main():
    """Command line interface for final document generator"""
    if len(sys.argv) < 2:
        print("Usage: python generate_final_document.py <document_folder> [format]")
        print("Formats: pdf, docx, both (default: both)")
        return
    
    document_folder = sys.argv[1]
    format_arg = sys.argv[2] if len(sys.argv) > 2 else 'both'
    
    # Parse format argument
    if format_arg.lower() in ['pdf', 'docx', 'both']:
        formats = [format_arg.lower()]
    else:
        print(f"Invalid format: {format_arg}. Use: pdf, docx, or both")
        return
    
    generator = FinalDocumentGenerator()
    
    try:
        generated_files = generator.generate_documents(document_folder, formats)
        
        print("\n" + "="*50)
        print("FINAL DOCUMENTS GENERATED")
        print("="*50)
        
        for format_type, file_path in generated_files.items():
            print(f"{format_type.upper()}: {file_path}")
        
        if not generated_files:
            print("No documents were generated successfully.")
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
