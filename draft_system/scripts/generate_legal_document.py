#!/usr/bin/env python3
"""
Legal Document Generator - Converts YAML to formatted Word documents
Based on the Legal Document Generation System

Usage: python generate_legal_document.py input.yml [output.docx]
"""

import yaml
import sys
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

class LegalDocumentGenerator:
    def __init__(self):
        self.current_paragraph_number = 1
        self.labels = {}  # Store label -> paragraph number mapping
        self.doc = None
        
    def generate_document(self, yaml_file, output_file=None):
        """Main entry point - generate Word document from YAML"""
        
        # Load and validate YAML
        try:
            with open(yaml_file, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
        except Exception as e:
            raise Exception(f"Error loading YAML file: {e}")
        
        # Validate required structure
        self._validate_yaml(data)
        
        # Create new document
        self.doc = Document()
        
        # Set up document formatting
        self._setup_document_formatting(data.get('metadata', {}))
        
        # Generate caption if present
        if 'caption' in data:
            self._create_caption(data['caption'], data.get('title', ''))
        
        # Two-pass processing for cross-references
        # Pass 1: Collect all labels
        self._collect_labels(data.get('body', []))
        
        # Pass 2: Generate content with resolved references
        self._generate_body(data.get('body', []), data.get('metadata', {}))
        
        # Add prayer for relief if present
        if 'prayer_for_relief' in data:
            self._generate_prayer_for_relief(data['prayer_for_relief'])
        
        # Add signature block if present
        if 'signature' in data:
            self._generate_signature_block(data['signature'])
        
        # Add certificate of service if present
        if 'certificate_of_service' in data:
            self._generate_certificate_of_service(data['certificate_of_service'])
        
        # Save document
        if output_file is None:
            output_file = yaml_file.replace('.yml', '.docx').replace('.yaml', '.docx')
        
        self.doc.save(output_file)
        print(f"Legal document generated: {output_file}")
        return output_file
    
    def _validate_yaml(self, data):
        """Validate YAML structure"""
        if 'body' not in data:
            raise Exception("YAML must contain 'body' section")
        
        # Check validation rules if present
        if 'validation' in data:
            validation = data['validation']
            if 'required_sections' in validation:
                self._check_required_sections(data['body'], validation['required_sections'])
    
    def _check_required_sections(self, body, required_sections):
        """Check that all required sections are present"""
        found_sections = []
        for section in body:
            if 'heading' in section:
                found_sections.append(section['heading'])
        
        missing = [req for req in required_sections if req not in found_sections]
        if missing:
            raise Exception(f"Missing required sections: {missing}")
    
    def _setup_document_formatting(self, metadata):
        """Set up document-wide formatting"""
        style_options = metadata.get('style_options', {})
        
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set default font and formatting
        style = self.doc.styles['Normal']
        font = style.font
        font.name = style_options.get('font_family', 'Times New Roman')
        font.size = Pt(style_options.get('font_size', 14))
        
        # Set paragraph formatting
        paragraph_format = style.paragraph_format
        if style_options.get('double_spaced', True):
            paragraph_format.line_spacing = 2.0
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.space_after = Pt(0)
    
    def _create_caption(self, caption_data, title):
        """Create legal document caption"""
        
        # Court identifier (centered, bold)
        if 'court_identifier' in caption_data and caption_data['court_identifier']:
            court_para = self.doc.add_paragraph()
            court_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            court_run = court_para.add_run(caption_data['court_identifier'])
            court_run.font.name = 'Times New Roman'
            court_run.font.size = Pt(14)
            court_run.bold = True
            court_para.paragraph_format.line_spacing = 2.0
        
        # Jurisdiction line if present
        if 'jurisdiction_line' in caption_data and caption_data['jurisdiction_line']:
            juris_para = self.doc.add_paragraph()
            juris_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            juris_run = juris_para.add_run(caption_data['jurisdiction_line'])
            juris_run.font.name = 'Times New Roman'
            juris_run.font.size = Pt(14)
            juris_run.bold = True
            juris_para.paragraph_format.line_spacing = 2.0
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Create caption table with parties and case info
        table = self.doc.add_table(rows=1, cols=2)
        table.style = None  # Remove default styling
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set table borders - only middle vertical line
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl = table._tbl
        tblPr = tbl.tblPr

        # Remove all borders first
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)

        # Add only the inside vertical border (middle line)
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'single')
        insideV.set(qn('w:sz'), '4')
        insideV.set(qn('w:space'), '0')
        insideV.set(qn('w:color'), '000000')
        tblBorders.append(insideV)

        tblPr.append(tblBorders)
        
        # Left cell - parties
        left_cell = table.cell(0, 0)
        left_cell.width = Inches(4.0)
        
        # Right cell - case info
        right_cell = table.cell(0, 1)
        right_cell.width = Inches(2.5)
        
        # Populate left cell with parties
        self._populate_parties_cell(left_cell, caption_data.get('parties', []))
        
        # Populate right cell with case info
        self._populate_case_info_cell(right_cell, caption_data)
        
        # Add spacing after caption
        self.doc.add_paragraph()
        
        # Add document title if present
        if title:
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title.upper())
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(14)
            title_run.bold = True
            title_para.paragraph_format.line_spacing = 2.0
            title_para.paragraph_format.space_after = Pt(12)
    
    def _populate_parties_cell(self, cell, parties):
        """Populate the left cell of caption with party information"""
        cell_para = cell.paragraphs[0]
        cell_para.paragraph_format.line_spacing = 2.0
        
        for i, party_group in enumerate(parties):
            role = party_group.get('role', '')
            names = party_group.get('names', [])
            
            # Add party names
            for name in names:
                run = cell_para.add_run(name)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                cell_para.add_run('\n')
            
            # Add role
            if role:
                cell_para.add_run('\n')
                role_run = cell_para.add_run(f"{role},")
                role_run.font.name = 'Times New Roman'
                role_run.font.size = Pt(14)
                cell_para.add_run('\n\n')
            
            # Add "v." between plaintiff and defendant
            if i == 0 and len(parties) > 1:
                v_run = cell_para.add_run("v.")
                v_run.font.name = 'Times New Roman'
                v_run.font.size = Pt(14)
                cell_para.add_run('\n\n')
    
    def _populate_case_info_cell(self, cell, caption_data):
        """Populate the right cell of caption with case information"""
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_para.paragraph_format.line_spacing = 2.0
        
        # Case number
        case_number = caption_data.get('case_number', '[TO BE ASSIGNED]')
        case_run = cell_para.add_run(f"Case No.\n\n{case_number}")
        case_run.font.name = 'Times New Roman'
        case_run.font.size = Pt(14)
        case_run.bold = True
        
        # Judge if present
        if 'judge' in caption_data and caption_data['judge']:
            cell_para.add_run(f"\n\n{caption_data['judge']}")
    
    def _collect_labels(self, body_sections):
        """First pass: collect all labels and their paragraph numbers"""
        temp_paragraph_number = 1
        
        for section in body_sections:
            temp_paragraph_number = self._collect_labels_from_section(section, temp_paragraph_number)
    
    def _collect_labels_from_section(self, section, paragraph_number):
        """Recursively collect labels from a section"""
        if 'content' in section:
            for content_item in section['content']:
                if content_item.get('type') == 'paragraph':
                    # Check for labels in paragraph data
                    data = content_item.get('data', '')
                    labels_in_para = re.findall(r'\{\{LABEL:([^}]+)\}\}', data)
                    for label in labels_in_para:
                        self.labels[label] = paragraph_number
                    paragraph_number += 1
                
                elif content_item.get('type') == 'section':
                    # Recursively process nested sections
                    paragraph_number = self._collect_labels_from_section(content_item, paragraph_number)
        
        return paragraph_number

    def _generate_body(self, body_sections, metadata):
        """Generate the main body content"""
        style_options = metadata.get('style_options', {})
        numbered_paragraphs = style_options.get('numbered_paragraphs', True)

        for section in body_sections:
            self._generate_section(section, numbered_paragraphs)

    def _generate_section(self, section, numbered_paragraphs, level=1):
        """Generate a section with heading and content"""

        # Add section heading
        if 'heading' in section:
            self._add_section_heading(section['heading'], level)

        # Process section content
        if 'content' in section:
            for content_item in section['content']:
                self._generate_content_item(content_item, numbered_paragraphs, level)

    def _add_section_heading(self, heading_text, level):
        """Add a section heading with appropriate formatting"""
        para = self.doc.add_paragraph()

        if level == 1:
            # Major sections: centered, bold, all caps
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(heading_text.upper())
            run.bold = True
        elif level == 2:
            # Subsections: centered, bold
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(heading_text)
            run.bold = True
        elif level == 3:
            # Level 3: centered, bold (like TILA-RESPA subtitle)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(heading_text)
            run.bold = True
        else:
            # Lower level: left-aligned, bold
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(heading_text)
            run.bold = True

        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _generate_content_item(self, content_item, numbered_paragraphs, level):
        """Generate individual content items"""
        content_type = content_item.get('type', 'paragraph')

        if content_type == 'paragraph':
            self._add_paragraph(content_item, numbered_paragraphs)
        elif content_type == 'list':
            self._add_list(content_item)
        elif content_type == 'section':
            # Nested section
            nested_level = content_item.get('heading_level', level + 1)
            self._generate_section(content_item, numbered_paragraphs, nested_level)

    def _add_paragraph(self, content_item, numbered_paragraphs):
        """Add a paragraph with optional numbering"""
        data = content_item.get('data', '')

        # Resolve cross-references
        data = self._resolve_references(data)

        # Remove label markers (they're invisible in final document)
        data = re.sub(r'\{\{LABEL:[^}]+\}\}', '', data)

        # Create paragraph
        para = self.doc.add_paragraph()
        para.paragraph_format.line_spacing = 2.0
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add paragraph number if enabled
        if numbered_paragraphs:
            # Set paragraph to start exactly at 1-inch margin with hanging indent
            para.paragraph_format.left_indent = Inches(0.25)  # Reduced from 0.5
            para.paragraph_format.first_line_indent = Inches(-0.25)

            # Add number
            num_run = para.add_run(f"{self.current_paragraph_number}. ")
            num_run.font.name = 'Times New Roman'
            num_run.font.size = Pt(14)
            num_run.bold = True

            self.current_paragraph_number += 1

        # Add paragraph text with formatting
        self._add_formatted_text(para, data)

    def _add_list(self, content_item):
        """Add a formatted list"""
        list_data = content_item.get('data', {})
        style = list_data.get('style', 'bullet')
        items = list_data.get('items', [])

        for i, item in enumerate(items):
            para = self.doc.add_paragraph()
            para.paragraph_format.line_spacing = 2.0
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Inches(1.0)
            para.paragraph_format.first_line_indent = Inches(-0.25)

            # Add list marker
            if style == 'lettered':
                marker = f"{chr(97 + i)}. "  # a., b., c., ...
            elif style == 'numbered':
                marker = f"{i + 1}. "  # 1., 2., 3., ...
            else:  # bullet
                marker = "• "

            marker_run = para.add_run(marker)
            marker_run.font.name = 'Times New Roman'
            marker_run.font.size = Pt(14)

            # Add item text with formatting
            self._add_formatted_text(para, item)

    def _add_formatted_text(self, paragraph, text):
        """Add text with legal formatting (bold statutes, etc.)"""
        # Split text by formatting markers
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text (legal citations)
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # Regular text
                run = paragraph.add_run(part)

            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    def _resolve_references(self, text):
        """Resolve cross-references in text"""
        # Handle range references: {{REF_RANGE:start:end}}
        def replace_range_ref(match):
            start_label = match.group(1)
            end_label = match.group(2)

            start_num = self.labels.get(start_label)
            end_num = self.labels.get(end_label)

            if start_num is not None and end_num is not None:
                if start_num == end_num:
                    return str(start_num)
                else:
                    return f"{start_num} through {end_num}"
            else:
                return f"[REF ERROR: {start_label}:{end_label}]"

        text = re.sub(r'\{\{REF_RANGE:([^:]+):([^}]+)\}\}', replace_range_ref, text)

        # Handle single references: {{REF:label}}
        def replace_single_ref(match):
            label = match.group(1)
            num = self.labels.get(label)
            return str(num) if num is not None else f"[REF ERROR: {label}]"

        text = re.sub(r'\{\{REF:([^}]+)\}\}', replace_single_ref, text)

        return text

    def _generate_prayer_for_relief(self, prayer_data):
        """Generate prayer for relief section"""
        # Add heading
        heading = prayer_data.get('heading', 'PRAYER FOR RELIEF')
        self._add_section_heading(heading, 1)

        # Add introduction
        intro_para = self.doc.add_paragraph()
        intro_para.paragraph_format.line_spacing = 2.0
        intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro_run = intro_para.add_run("WHEREFORE, Plaintiff respectfully requests that this Court enter judgment in favor of Plaintiff and against Defendants and award the following relief:")
        intro_run.font.name = 'Times New Roman'
        intro_run.font.size = Pt(14)

        # Add requests
        requests = prayer_data.get('requests', [])
        for i, request in enumerate(requests):
            para = self.doc.add_paragraph()
            para.paragraph_format.line_spacing = 2.0
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)

            # Add letter marker
            marker = f"{chr(65 + i)}. "  # A., B., C., ...
            marker_run = para.add_run(marker)
            marker_run.font.name = 'Times New Roman'
            marker_run.font.size = Pt(14)
            marker_run.bold = True

            # Add request text
            description = request.get('description', '')
            self._add_formatted_text(para, description)

        # Add standard requests
        standard_requests = []
        if prayer_data.get('attorney_fees', False):
            standard_requests.append("Reasonable attorney's fees and costs as provided by law;")
        if prayer_data.get('costs', False):
            standard_requests.append("Costs of suit;")
        if prayer_data.get('other_relief', False):
            standard_requests.append("Such other and further relief as this Court deems just and proper.")

        for request_text in standard_requests:
            para = self.doc.add_paragraph()
            para.paragraph_format.line_spacing = 2.0
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)

            # Continue lettering
            marker = f"{chr(65 + len(requests) + len(standard_requests) - len(standard_requests) + standard_requests.index(request_text))}. "
            marker_run = para.add_run(marker)
            marker_run.font.name = 'Times New Roman'
            marker_run.font.size = Pt(14)
            marker_run.bold = True

            text_run = para.add_run(request_text)
            text_run.font.name = 'Times New Roman'
            text_run.font.size = Pt(14)

    def _generate_signature_block(self, signature_data):
        """Generate signature block"""
        # Add spacing before signature
        self.doc.add_paragraph()

        # Closing
        closing = signature_data.get('closing', 'Respectfully submitted,')
        closing_para = self.doc.add_paragraph()
        closing_para.paragraph_format.line_spacing = 1.0  # Single spacing for signature block
        closing_run = closing_para.add_run(closing)
        closing_run.font.name = 'Times New Roman'
        closing_run.font.size = Pt(14)

        # Signature line
        self.doc.add_paragraph()
        sig_para = self.doc.add_paragraph()
        sig_para.paragraph_format.line_spacing = 1.0  # Single spacing
        sig_run = sig_para.add_run("_" * 40)
        sig_run.font.name = 'Times New Roman'
        sig_run.font.size = Pt(14)

        # Attorney information
        info_lines = []
        if 'name' in signature_data:
            info_lines.append(signature_data['name'])
        if 'bar_number' in signature_data:
            info_lines.append(signature_data['bar_number'])
        if 'title' in signature_data:
            info_lines.append(signature_data['title'])
        if 'firm_name' in signature_data:
            info_lines.append(signature_data['firm_name'])
        if 'address' in signature_data:
            info_lines.append(signature_data['address'])
        if 'phone' in signature_data:
            info_lines.append(signature_data['phone'])
        if 'email' in signature_data:
            info_lines.append(signature_data['email'])

        for line in info_lines:
            info_para = self.doc.add_paragraph()
            info_para.paragraph_format.line_spacing = 1.0  # Single spacing for signature block
            info_run = info_para.add_run(line)
            info_run.font.name = 'Times New Roman'
            info_run.font.size = Pt(14)

    def _generate_certificate_of_service(self, cert_data):
        """Generate certificate of service"""
        # Add page break before certificate of service
        self.doc.add_page_break()

        # Add heading
        heading = cert_data.get('heading', 'CERTIFICATE OF SERVICE')
        self._add_section_heading(heading, 1)

        # Add certificate text
        cert_text = cert_data.get('text', '')
        if cert_text:
            para = self.doc.add_paragraph()
            para.paragraph_format.line_spacing = 2.0
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._add_formatted_text(para, cert_text)

        # Add service details
        if 'served_parties' in cert_data:
            for party in cert_data['served_parties']:
                para = self.doc.add_paragraph()
                para.paragraph_format.line_spacing = 2.0
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Inches(0.5)

                name = party.get('name', '')
                method = party.get('method', '')
                text = f"{name}: {method}"

                text_run = para.add_run(text)
                text_run.font.name = 'Times New Roman'
                text_run.font.size = Pt(14)

        # Add date and signature
        if 'service_date' in cert_data:
            date_para = self.doc.add_paragraph()
            date_para.paragraph_format.line_spacing = 2.0
            date_run = date_para.add_run(f"This {cert_data['service_date']}.")
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(14)

        # Signature line for certificate
        self.doc.add_paragraph()
        cert_sig_para = self.doc.add_paragraph()
        cert_sig_para.paragraph_format.line_spacing = 2.0
        cert_sig_run = cert_sig_para.add_run("_" * 40)
        cert_sig_run.font.name = 'Times New Roman'
        cert_sig_run.font.size = Pt(14)

        if 'served_by' in cert_data:
            served_by_para = self.doc.add_paragraph()
            served_by_para.paragraph_format.line_spacing = 2.0
            served_by_run = served_by_para.add_run(cert_data['served_by'])
            served_by_run.font.name = 'Times New Roman'
            served_by_run.font.size = Pt(14)


def main():
    """Main execution function"""
    if len(sys.argv) < 2:
        print("Usage: python generate_legal_document.py input.yml [output.docx]")
        sys.exit(1)

    yaml_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(yaml_file):
        print(f"Error: YAML file '{yaml_file}' not found")
        sys.exit(1)

    try:
        generator = LegalDocumentGenerator()
        output_path = generator.generate_document(yaml_file, output_file)
        print(f"Success! Legal document generated: {output_path}")

    except Exception as e:
        print(f"Error generating document: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
